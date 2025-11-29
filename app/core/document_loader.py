# app/core/document_loader.py
import io
from typing import Union
import fitz  # PyMuPDF
import docx

def load_pdf_bytes(content: bytes) -> str:
    """Extract text from PDF bytes (uploaded file)."""
    doc = fitz.open(stream=content, filetype="pdf")
    texts = []
    for page in doc:
        page_text = page.get_text()
        if page_text:
            texts.append(page_text)
    return "\n".join(texts)

def load_pdf_path(path: str) -> str:
    with fitz.open(path) as doc:
        return "\n".join(page.get_text() for page in doc)

def load_docx_path(path: str) -> str:
    doc = docx.Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text]
    return "\n".join(paragraphs)

def load_docx_bytes(content: bytes) -> str:
    bio = io.BytesIO(content)
    doc = docx.Document(bio)
    paragraphs = [p.text for p in doc.paragraphs if p.text]
    return "\n".join(paragraphs)

def load_txt_bytes(content: bytes) -> str:
    return content.decode("utf-8", errors="ignore")

def extract_text_from_file(filename: str, content: bytes) -> str:
    """
    Generic loader used by ingestion endpoint.
    filename: original filename (used to decide parser)
    content: raw bytes of uploaded file
    """
    name = filename.lower()
    if name.endswith(".pdf"):
        return load_pdf_bytes(content)
    if name.endswith(".docx"):
        return load_docx_bytes(content)
    if name.endswith(".txt"):
        return load_txt_bytes(content)
    # fallback try to decode
    try:
        return content.decode("utf-8", errors="ignore")
    except Exception:
        return ""
